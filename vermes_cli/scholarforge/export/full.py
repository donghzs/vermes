"""
ScholarForge 导出模块 — 统一导出 PDF / Word / LaTeX / Markdown / BibTeX
- PDF: Markdown → HTML → weasyprint（纯 Python，无需 LaTeX）
- Word: Markdown → pandoc → docx
- LaTeX / Markdown / BibTeX: 复用原有 export 模块
"""
import io
import os
import re
import html
import logging
import subprocess
from typing import Optional

logger = logging.getLogger("scholarforge.export.full")


# ============================================================================
# Markdown 渲染
# ============================================================================
def _markdown_to_html(md_text: str) -> str:
    """Markdown → HTML（学术论文样式，干净、无特殊符号）"""
    from markdown_it import MarkdownIt
    md = MarkdownIt("commonmark", {"html": True, "breaks": False, "linkify": True})
    md.enable(["table", "strikethrough"])
    return md.render(md_text)


def _latex_to_md_text(content: str) -> str:
    """若 content 是 LaTeX 格式（含 \\section 等），先粗略转 Markdown。"""
    # 如果含 \\section 或 \\begin{document}，视为 LaTeX
    if "\\section" in content or "\\begin{document}" in content:
        # 章节
        content = re.sub(r"\\section\{([^}]+)\}", r"## \1", content)
        content = re.sub(r"\\subsection\{([^}]+)\}", r"### \1", content)
        content = re.sub(r"\\subsubsection\{([^}]+)\}", r"#### \1", content)
        # 强调
        content = re.sub(r"\\textbf\{([^}]+)\}", r"**\1**", content)
        content = re.sub(r"\\textit\{([^}]+)\}", r"*\1*", content)
        content = re.sub(r"\\emph\{([^}]+)\}", r"*\1*", content)
        # 数学
        content = re.sub(r"\$([^$]+)\$", r"$\1$", content)
        # 清理命令
        content = re.sub(r"\\begin\{[^}]+\}", "", content)
        content = re.sub(r"\\end\{[^}]+\}", "", content)
        content = re.sub(r"\\\\", "\n\n", content)
    return content


# ============================================================================
# 标题/作者/参考文献预处理
# ============================================================================
def _build_paper_text(title: str, content: str, papers: list, abstract: str = "") -> str:
    """组装完整论文 Markdown 文本
    
    智能检测：如果 content 已包含标题/摘要/参考文献，不再重复添加。
    WritingAgent._assemble_full_paper 已将完整结构写入 ctx.draft，
    本函数只做兜底补充。
    """
    text = _latex_to_md_text(content or "")
    parts = []
    
    # 标题：仅当 content 不以 Markdown 标题开头时才添加
    if not text.strip().startswith("# "):
        parts.append(f"# {title}")
        parts.append("")
    
    # 摘要：仅当 content 不含摘要时才添加
    if abstract and re.search(r'(?i)#{1,3}\s*(摘要|abstract)', text) is None:
        parts.append("## 摘要")
        parts.append("")
        parts.append(abstract.strip())
        parts.append("")
    
    # 主体内容
    if text.strip():
        parts.append(text.strip())
        parts.append("")
    
    # 参考文献：仅当 content 不含参考文献节时才添加
    if papers and re.search(r'(?i)#{1,3}\s*(参考文献|references)', text) is None:
        parts.append("## 参考文献")
        parts.append("")
        for i, p in enumerate(papers, 1):
            p_dict = p.to_dict() if hasattr(p, "to_dict") else (
                p if isinstance(p, dict) else {
                    "title": getattr(p, "title", ""),
                    "authors": getattr(p, "authors", []),
                    "year": getattr(p, "year", ""),
                    "venue": getattr(p, "venue", ""),
                }
            )
            authors = p_dict.get("authors", []) or []
            if isinstance(authors, list):
                authors_str = ", ".join(authors[:3]) + (" et al." if len(authors) > 3 else "")
            else:
                authors_str = str(authors)
            title_text = p_dict.get("title", "").strip()
            year = p_dict.get("year", "")
            venue = p_dict.get("venue", "")
            parts.append(f"[{i}] {authors_str}. **{title_text}**. {venue} ({year}).")
        parts.append("")

    return "\n".join(parts)


# ============================================================================
# PDF 导出（Markdown → HTML → WeasyPrint）
# ============================================================================
from .pdf_css import _PDF_CSS



def export_pdf(title: str, content: str, papers: list, abstract: str = "") -> bytes:
    """导出 PDF — Markdown → HTML → WeasyPrint"""
    md_text = _build_paper_text(title, content, papers, abstract)
    html_body = _markdown_to_html(md_text)

    # 中文换行优化
    html_body = re.sub(
        r"(<p[^>]*>)(.*?)(</p>)",
        lambda m: m.group(1) + m.group(2) + m.group(3),
        html_body,
        flags=re.DOTALL,
    )

    full_html = (
        "<!DOCTYPE html>\n"
        '<html lang="zh-CN">\n'
        "<head>\n"
        '<meta charset="UTF-8">\n'
        f"<title>{html.escape(title)}</title>\n"
        f"<style>{_PDF_CSS}</style>\n"
        "</head>\n"
        f"<body>\n{html_body}\n</body>\n"
        "</html>\n"
    )

    try:
        # macOS Homebrew: weasyprint 需要 DYLD_FALLBACK_LIBRARY_PATH 才能找到 gobject
        import os as _os
        if _os.path.exists("/opt/homebrew/lib/libgobject-2.0.dylib"):
            _os.environ.setdefault("DYLD_FALLBACK_LIBRARY_PATH", "/opt/homebrew/lib")
        from weasyprint import HTML
        buf = io.BytesIO()
        HTML(string=full_html, base_url=".").write_pdf(target=buf)
        return buf.getvalue()
    except Exception as e:
        logger.warning(f"WeasyPrint PDF export failed: {e}, trying reportlab fallback...")
        # Fallback: reportlab (纯 Python，无系统库依赖)
        try:
            return _export_pdf_reportlab(title, md_text)
        except Exception as e2:
            logger.error(f"reportlab PDF fallback also failed: {e2}", exc_info=True)
            # Final fallback: 返回带打印脚本的 HTML 页面
            logger.info("Returning HTML for browser print-to-PDF")
            print_html = (
                "<!DOCTYPE html>\n<html lang=\"zh-CN\">\n<head>\n"
                '<meta charset=\"UTF-8\">\n'
                f"<title>{html.escape(title)}</title>\n"
                '<style>body{font-family:sans-serif;padding:2em;max-width:800px;margin:0 auto}</style>\n'
                '</head>\n<body>\n'
                '<h1 style="color:#999">⚠️ PDF 导出失败，请使用浏览器打印功能（Ctrl/Cmd+P）保存为 PDF</h1>\n'
                '<hr/>\n'
                f'{html_body}\n'
                '<script>window.onload=()=>{window.print()}</script>\n'
                '</body>\n</html>\n'
            )
            return print_html.encode("utf-8")


def _export_pdf_reportlab(title: str, md_text: str) -> bytes:
    """reportlab fallback PDF 导出 — 纯 Python，无系统库依赖
    
    简单排版：标题 + 正文段落，不渲染复杂 HTML 样式。
    """
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    import io as _io

    # 尝试注册中文字体
    font_name = "Helvetica"
    for font_path in [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont("CJKFont", font_path, subfontIndex=0))
                font_name = "CJKFont"
                break
            except Exception:
                continue

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4,
                            topMargin=2.5*cm, bottomMargin=2.5*cm,
                            leftMargin=2.5*cm, rightMargin=2.5*cm)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("CustomTitle", parent=styles["Title"],
                                 fontName=font_name, fontSize=18, leading=24,
                                 alignment=TA_CENTER, spaceAfter=20)
    body_style = ParagraphStyle("CustomBody", parent=styles["Normal"],
                                fontName=font_name, fontSize=11, leading=18,
                                alignment=TA_LEFT, spaceAfter=6, firstLineIndent=24)
    h2_style = ParagraphStyle("CustomH2", parent=styles["Heading2"],
                              fontName=font_name, fontSize=14, leading=20,
                              spaceBefore=12, spaceAfter=6)

    story = []
    story.append(Paragraph(_escape_xml(title), title_style))
    story.append(Spacer(1, 0.5*cm))

    for line in md_text.split("\n"):
        stripped = line.strip()
        if not stripped:
            continue
        # 简单 Markdown 转换
        if stripped.startswith("### "):
            story.append(Paragraph(f'<b>{_escape_xml(stripped[4:])}</b>', h2_style))
        elif stripped.startswith("## "):
            story.append(Paragraph(f'<b>{_escape_xml(stripped[3:])}</b>', h2_style))
        elif stripped.startswith("# "):
            story.append(Paragraph(f'<b>{_escape_xml(stripped[2:])}</b>', title_style))
        else:
            # 转义 XML 特殊字符
            safe = _escape_xml(stripped)
            story.append(Paragraph(safe, body_style))

    doc.build(story)
    return buf.getvalue()


def _escape_xml(text: str) -> str:
    """转义 XML/HTML 特殊字符"""
    return (text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;"))


# ============================================================================
# Word 导出（Markdown → pandoc → docx）
# ============================================================================
def export_docx(title: str, content: str, papers: list, abstract: str = "") -> bytes:
    """导出 Word (.docx) — Markdown → pandoc"""
    md_text = _build_paper_text(title, content, papers, abstract)

    # 准备 pandoc 元数据
    metadata = [
        "---",
        f"title: {title}",
        "mainfont: Noto Serif CJK SC",
        "sansfont: Noto Sans CJK SC",
        "monofont: DejaVu Sans Mono",
        "geometry: margin=2.5cm",
        "---",
    ]
    full_md = "\n".join(metadata) + "\n\n" + md_text

    try:
        # 尝试 pandoc
        result = subprocess.run(
            ["pandoc", "-f", "markdown", "-t", "docx", "--standalone"],
            input=full_md.encode("utf-8"),
            capture_output=True,
            timeout=30,
        )
        if result.returncode == 0 and result.stdout:
            # 后处理：把 pandoc 生成的默认有框表格改为学术三线表
            try:
                import io as _io
                from docx import Document
                d = Document(_io.BytesIO(result.stdout))
                _style_three_line_tables(d)
                out = _io.BytesIO()
                d.save(out)
                return out.getvalue()
            except Exception as e:
                logger.warning(f"docx 三线表后处理失败，回退原始 pandoc 输出: {e}")
                return result.stdout
        # 失败时回退到纯 Python
        logger.warning(f"pandoc failed: {result.stderr.decode()}, falling back to python-docx")
    except FileNotFoundError:
        logger.warning("pandoc not found, falling back to python-docx")
    except Exception as e:
        logger.warning(f"pandoc error: {e}, falling back to python-docx")

    # 回退方案：python-docx
    return _export_docx_fallback(title, content, papers, abstract, md_text)


def _export_docx_fallback(title: str, content: str, papers: list, abstract: str, md_text: str) -> bytes:
    """python-docx 回退方案（无 pandoc 时使用）。表格渲染为学术三线表。"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise RuntimeError("需要安装 pandoc 或 python-docx")

    doc = Document()
    # 标题样式
    title_para = doc.add_heading(title, level=0)

    # 简单按行解析 Markdown（表格行先缓冲，连续 | 行合并为一个表格）
    table_lines: list[str] = []

    def _flush_table():
        if table_lines:
            _add_markdown_table(doc, table_lines)
            table_lines.clear()

    for line in md_text.split("\n"):
        line = line.rstrip()
        if not line:
            _flush_table()
            continue
        if line.startswith("# "):
            _flush_table()
            continue  # 标题已添加
        elif line.startswith("## "):
            _flush_table()
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            _flush_table()
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            _flush_table()
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("---"):
            _flush_table()
            doc.add_paragraph("─" * 30)
        elif line.startswith("```"):
            _flush_table()
            continue  # 简化：跳过代码块标记
        elif re.match(r"^\d+\.\s", line):
            _flush_table()
            doc.add_paragraph(line, style="List Number")
        elif line.startswith("- "):
            _flush_table()
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("|"):
            table_lines.append(line)
        else:
            _flush_table()
            para = doc.add_paragraph()
            _add_runs(para, line)
    _flush_table()

    # 全部表格改为学术三线表
    _style_three_line_tables(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_markdown_table(doc, lines: list[str]):
    """把一组 markdown 表格行解析为 python-docx 表格（首行视为表头，分隔行 --- 跳过）。"""
    rows: list[list[str]] = []
    for ln in lines:
        cells = [c.strip() for c in ln.strip().strip("|").split("|")]
        if cells and all(re.match(r"^:?-+:?$", c) for c in cells):
            continue  # 分隔行
        rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    for ri, cells in enumerate(rows):
        tcells = table.rows[ri].cells
        for ci in range(ncols):
            text = cells[ci] if ci < len(cells) else ""
            para = tcells[ci].paragraphs[0]
            para.text = ""
            _add_runs(para, text)


def _set_cell_border(cell, edge: str, sz: int, val: str = "single", color: str = "auto"):
    """设置单元格某条边的边框（用于三线表表头下横线）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    el = tcBorders.find(qn(f"w:{edge}"))
    if el is None:
        el = OxmlElement(f"w:{edge}")
        tcBorders.append(el)
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), str(sz))
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)


def _style_three_line_tables(doc):
    """把文档中所有表格改为学术三线表：上下粗线 1.5pt(24)、表头下细线 0.75pt(12)、无竖线、无内部横线。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.tblPr
        borders = tblPr.find(qn("w:tblBorders"))
        if borders is None:
            borders = OxmlElement("w:tblBorders")
            tblPr.append(borders)
        for edge, sz in (
            ("top", 24), ("bottom", 24), ("left", 0),
            ("right", 0), ("insideV", 0), ("insideH", 0),
        ):
            el = borders.find(qn(f"w:{edge}"))
            if el is None:
                el = OxmlElement(f"w:{edge}")
                borders.append(el)
            if sz == 0:
                el.set(qn("w:val"), "none")
            else:
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), str(sz))
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "auto")
        # 表头（首行）下方细线
        if table.rows:
            for cell in table.rows[0].cells:
                _set_cell_border(cell, "bottom", 12)


def _add_runs(para, text: str):
    """向段落添加 runs（处理粗体/斜体）"""
    from docx import Document
    # 拆分 **bold** / *italic* / `code`
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            r = para.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("`"):
            r = para.add_run(token[1:-1])
            r.font.name = "Courier New"
        else:
            r = para.add_run(token[1:-1])
            r.italic = True
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


# ============================================================================
# LaTeX 导出（已被原 export.latex 实现）
# ============================================================================
def export_latex(title: str, content: str, papers: list, template: str = "ieee", abstract: str = "") -> str:
    """导出 LaTeX"""
    from .latex import format_export_latex
    return format_export_latex(title, content, papers, template=template)


# ============================================================================
# Markdown / BibTeX 导出（已被原 export 实现）
# ============================================================================
def export_markdown(title: str, content: str, papers: list, abstract: str = "") -> str:
    from . import format_export_markdown
    return format_export_markdown(title, content, papers)


def export_bibtex(papers: list) -> str:
    from . import format_export_bibtex
    return format_export_bibtex(papers)


def export_csl_json(papers: list) -> str:
    """导出 Zotero 兼容的 CSL JSON（不读 SQLite，仅基于传入文献列表）。"""
    from . import format_export_csl_json
    return format_export_csl_json(papers)
