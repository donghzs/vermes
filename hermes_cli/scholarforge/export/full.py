"""
ScholarForge 导出模块 — 统一导出 PDF / Word / LaTeX / Markdown / BibTeX
- PDF: Markdown → HTML → weasyprint（纯 Python，无需 LaTeX）
- Word: Markdown → pandoc → docx
- LaTeX / Markdown / BibTeX: 复用原有 export 模块
"""
import io
import os
import re
import html
import logging
import subprocess
from typing import Optional

logger = logging.getLogger("scholarforge.export.full")


# ============================================================================
# Markdown 渲染
# ============================================================================
def _markdown_to_html(md_text: str) -> str:
    """Markdown → HTML（学术论文样式，干净、无特殊符号）"""
    from markdown_it import MarkdownIt
    md = MarkdownIt("commonmark", {"html": True, "breaks": False, "linkify": True})
    md.enable(["table", "strikethrough"])
    return md.render(md_text)


def _latex_to_md_text(content: str) -> str:
    """若 content 是 LaTeX 格式（含 \\section 等），先粗略转 Markdown。"""
    # 如果含 \\section 或 \\begin{document}，视为 LaTeX
    if "\\section" in content or "\\begin{document}" in content:
        # 章节
        content = re.sub(r"\\section\{([^}]+)\}", r"## \1", content)
        content = re.sub(r"\\subsection\{([^}]+)\}", r"### \1", content)
        content = re.sub(r"\\subsubsection\{([^}]+)\}", r"#### \1", content)
        # 强调
        content = re.sub(r"\\textbf\{([^}]+)\}", r"**\1**", content)
        content = re.sub(r"\\textit\{([^}]+)\}", r"*\1*", content)
        content = re.sub(r"\\emph\{([^}]+)\}", r"*\1*", content)
        # 数学
        content = re.sub(r"\$([^$]+)\$", r"$\1$", content)
        # 清理命令
        content = re.sub(r"\\begin\{[^}]+\}", "", content)
        content = re.sub(r"\\end\{[^}]+\}", "", content)
        content = re.sub(r"\\\\", "\n\n", content)
    return content


# ============================================================================
# 标题/作者/参考文献预处理
# ============================================================================
def _build_paper_text(title: str, content: str, papers: list, abstract: str = "") -> str:
    """组装完整论文 Markdown 文本
    
    智能检测：如果 content 已包含标题/摘要/参考文献，不再重复添加。
    WritingAgent._assemble_full_paper 已将完整结构写入 ctx.draft，
    本函数只做兜底补充。
    """
    text = _latex_to_md_text(content or "")
    parts = []
    
    # 标题：仅当 content 不以 Markdown 标题开头时才添加
    if not text.strip().startswith("# "):
        parts.append(f"# {title}")
        parts.append("")
    
    # 摘要：仅当 content 不含摘要时才添加
    if abstract and re.search(r'(?i)#{1,3}\s*(摘要|abstract)', text) is None:
        parts.append("## 摘要")
        parts.append("")
        parts.append(abstract.strip())
        parts.append("")
    
    # 主体内容
    if text.strip():
        parts.append(text.strip())
        parts.append("")
    
    # 参考文献：仅当 content 不含参考文献节时才添加
    if papers and re.search(r'(?i)#{1,3}\s*(参考文献|references)', text) is None:
        parts.append("## 参考文献")
        parts.append("")
        for i, p in enumerate(papers, 1):
            p_dict = p.to_dict() if hasattr(p, "to_dict") else (
                p if isinstance(p, dict) else {
                    "title": getattr(p, "title", ""),
                    "authors": getattr(p, "authors", []),
                    "year": getattr(p, "year", ""),
                    "venue": getattr(p, "venue", ""),
                }
            )
            authors = p_dict.get("authors", []) or []
            if isinstance(authors, list):
                authors_str = ", ".join(authors[:3]) + (" et al." if len(authors) > 3 else "")
            else:
                authors_str = str(authors)
            title_text = p_dict.get("title", "").strip()
            year = p_dict.get("year", "")
            venue = p_dict.get("venue", "")
            parts.append(f"[{i}] {authors_str}. **{title_text}**. {venue} ({year}).")
        parts.append("")

    return "\n".join(parts)


# ============================================================================
# PDF 导出（Markdown → HTML → WeasyPrint）
# ============================================================================
_PDF_CSS = """
@page {
    size: A4;
    margin: 2.5cm 2cm 2.5cm 2cm;
    @bottom-center {
        content: counter(page) " / " counter(pages);
        font-family: "PingFang SC", "Heiti SC", "DejaVu Sans", sans-serif;
        font-size: 9pt;
        color: #666;
    }
    @top-right {
        content: "ScholarForge";
        font-family: "PingFang SC", "Heiti SC", "DejaVu Sans", sans-serif;
        font-size: 8pt;
        color: #999;
    }
}
body {
    font-family: "PingFang SC", "Heiti SC", "STSong", "STKaiti", serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #222;
    text-align: justify;
}
h1 {
    font-size: 22pt;
    text-align: center;
    margin: 1.5em 0 1em 0;
    font-weight: bold;
    color: #1a1a1a;
    border-bottom: 2px solid #333;
    padding-bottom: 0.3em;
}
h2 {
    font-size: 15pt;
    margin: 1.5em 0 0.8em 0;
    font-weight: bold;
    color: #2a2a2a;
    border-bottom: 1px solid #ddd;
    padding-bottom: 0.2em;
}
h3 {
    font-size: 12.5pt;
    margin: 1.2em 0 0.6em 0;
    font-weight: bold;
    color: #333;
}
h4 {
    font-size: 11.5pt;
    margin: 1em 0 0.5em 0;
    font-weight: bold;
}
p {
    margin: 0.5em 0;
    text-indent: 0;
}
ul, ol {
    margin: 0.5em 0;
    padding-left: 2em;
}
li {
    margin: 0.3em 0;
}
strong { font-weight: bold; }
em { font-style: italic; }
code {
    font-family: "DejaVu Sans Mono", "Courier New", monospace;
    font-size: 10pt;
    background: #f5f5f5;
    padding: 0.1em 0.3em;
    border-radius: 3px;
}
pre {
    background: #f5f5f5;
    padding: 0.8em;
    border-radius: 4px;
    overflow-x: auto;
    font-size: 9.5pt;
    line-height: 1.4;
}
pre code { background: transparent; padding: 0; }
blockquote {
    border-left: 3px solid #ccc;
    margin: 0.8em 0;
    padding: 0.3em 1em;
    color: #555;
    background: #fafafa;
}
table {
    border-collapse: collapse;
    margin: 1em 0;
    width: 100%;
    font-size: 10pt;
}
th, td {
    border: 1px solid #999;
    padding: 0.4em 0.7em;
    text-align: left;
}
th {
    background: #eee;
    font-weight: bold;
}
hr {
    border: none;
    border-top: 1px solid #ddd;
    margin: 1.5em 0;
}
a { color: #1a5490; text-decoration: none; }
a:hover { text-decoration: underline; }
"""


def export_pdf(title: str, content: str, papers: list, abstract: str = "") -> bytes:
    """导出 PDF — Markdown → HTML → WeasyPrint"""
    md_text = _build_paper_text(title, content, papers, abstract)
    html_body = _markdown_to_html(md_text)

    # 中文换行优化
    html_body = re.sub(
        r"(<p[^>]*>)(.*?)(</p>)",
        lambda m: m.group(1) + m.group(2) + m.group(3),
        html_body,
        flags=re.DOTALL,
    )

    full_html = (
        "<!DOCTYPE html>\n"
        '<html lang="zh-CN">\n'
        "<head>\n"
        '<meta charset="UTF-8">\n'
        f"<title>{html.escape(title)}</title>\n"
        f"<style>{_PDF_CSS}</style>\n"
        "</head>\n"
        f"<body>\n{html_body}\n</body>\n"
        "</html>\n"
    )

    try:
        # macOS Homebrew: weasyprint 需要 DYLD_FALLBACK_LIBRARY_PATH 才能找到 gobject
        import os as _os
        if _os.path.exists("/opt/homebrew/lib/libgobject-2.0.dylib"):
            _os.environ.setdefault("DYLD_FALLBACK_LIBRARY_PATH", "/opt/homebrew/lib")
        from weasyprint import HTML
        buf = io.BytesIO()
        HTML(string=full_html, base_url=".").write_pdf(target=buf)
        return buf.getvalue()
    except Exception as e:
        logger.error(f"PDF export failed: {e}", exc_info=True)
        raise


# ============================================================================
# Word 导出（Markdown → pandoc → docx）
# ============================================================================
def export_docx(title: str, content: str, papers: list, abstract: str = "") -> bytes:
    """导出 Word (.docx) — Markdown → pandoc"""
    md_text = _build_paper_text(title, content, papers, abstract)

    # 准备 pandoc 元数据
    metadata = [
        "---",
        f"title: {title}",
        "mainfont: Noto Serif CJK SC",
        "sansfont: Noto Sans CJK SC",
        "monofont: DejaVu Sans Mono",
        "geometry: margin=2.5cm",
        "---",
    ]
    full_md = "\n".join(metadata) + "\n\n" + md_text

    try:
        # 尝试 pandoc
        result = subprocess.run(
            ["pandoc", "-f", "markdown", "-t", "docx", "--standalone"],
            input=full_md.encode("utf-8"),
            capture_output=True,
            timeout=30,
        )
        if result.returncode == 0 and result.stdout:
            return result.stdout
        # 失败时回退到纯 Python
        logger.warning(f"pandoc failed: {result.stderr.decode()}, falling back to python-docx")
    except FileNotFoundError:
        logger.warning("pandoc not found, falling back to python-docx")
    except Exception as e:
        logger.warning(f"pandoc error: {e}, falling back to python-docx")

    # 回退方案：python-docx
    return _export_docx_fallback(title, content, papers, abstract, md_text)


def _export_docx_fallback(title: str, content: str, papers: list, abstract: str, md_text: str) -> bytes:
    """python-docx 回退方案"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise RuntimeError("需要安装 pandoc 或 python-docx")

    doc = Document()
    # 标题样式
    title_para = doc.add_heading(title, level=0)

    # 简单按行解析 Markdown
    for line in md_text.split("\n"):
        line = line.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            continue  # 标题已添加
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("---"):
            doc.add_paragraph("─" * 30)
        elif line.startswith("```"):
            continue  # 简化：跳过代码块标记
        elif re.match(r"^\d+\.\s", line):
            doc.add_paragraph(line, style="List Number")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.startswith("|"):
            # 表格行 - 简单拼接
            doc.add_paragraph(line)
        else:
            # 普通段落：粗体/斜体
            para = doc.add_paragraph()
            _add_runs(para, line)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_runs(para, text: str):
    """向段落添加 runs（处理粗体/斜体）"""
    from docx import Document
    # 拆分 **bold** / *italic* / `code`
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")
    last = 0
    for m in pattern.finditer(text):
        if m.start() > last:
            para.add_run(text[last:m.start()])
        token = m.group(0)
        if token.startswith("**"):
            r = para.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("`"):
            r = para.add_run(token[1:-1])
            r.font.name = "Courier New"
        else:
            r = para.add_run(token[1:-1])
            r.italic = True
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


# ============================================================================
# LaTeX 导出（已被原 export.latex 实现）
# ============================================================================
def export_latex(title: str, content: str, papers: list, template: str = "ieee", abstract: str = "") -> str:
    """导出 LaTeX"""
    from .latex import format_export_latex
    return format_export_latex(title, content, papers, template=template)


# ============================================================================
# Markdown / BibTeX 导出（已被原 export 实现）
# ============================================================================
def export_markdown(title: str, content: str, papers: list, abstract: str = "") -> str:
    from . import format_export_markdown
    return format_export_markdown(title, content, papers)


def export_bibtex(papers: list) -> str:
    from . import format_export_bibtex
    return format_export_bibtex(papers)
