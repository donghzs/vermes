# -*- coding: utf-8 -*-
"""P4-4 T1 协同编辑后端：docx 段落提取/回存 + pdf pandoc 重生成 + xlsx 字节回存的真实往返测试。

docx 走 python-docx（环境已带 1.2.0），xlsx 在方案 A 由前端 SheetJS 管理、后端仅收字节。
pdf 重生成需 pandoc + 一个可用 PDF 引擎（LaTeX / weasyprint）：本沙箱若无可用引擎，pdf 用例自动 skip
（机制仍正确，待用户 Mac / DMG 运行时验证——那里可有 MacTeX 或 .venv 内 weasyprint）。
"""
import sys
import uuid
import os

sys.path.insert(0, "/Users/dongzusheng/Projects/vermes-electron")

import pytest


def _tmp_name(ext: str) -> str:
    return f"p4t1_{uuid.uuid4().hex}{ext}"


def _pdf_engine_available() -> bool:
    """探测本机是否有可用的 pandoc PDF 引擎（LaTeX 或 weasyprint 可 import）。"""
    import shutil
    import subprocess
    import tempfile

    pandoc = shutil.which('pandoc')
    if not pandoc:
        return False
    engine = None
    if shutil.which('weasyprint'):
        try:
            import weasyprint  # noqa: F401
            engine = '--pdf-engine=weasyprint'
        except Exception:
            engine = None
    with tempfile.NamedTemporaryFile('w', suffix='.md', delete=False, encoding='utf-8') as f:
        f.write('# t\n')
        mp = f.name
    out = tempfile.mktemp(suffix='.pdf')
    try:
        cmd = [pandoc, '-f', 'markdown', '-t', 'pdf'] + ([engine] if engine else []) + ['-o', out, mp]
        r = subprocess.run(cmd, capture_output=True, text=True, timeout=60)
        ok = r.returncode == 0 and os.path.exists(out) and open(out, 'rb').read(5).startswith(b'%PDF')
        return ok
    except Exception:
        return False
    finally:
        for p in (mp, out):
            if os.path.exists(p):
                os.unlink(p)


def test_apply_docx_paragraphs_roundtrip():
    """真实 docx：改第 2 段 → 回存 → 重读确认改了且其余段不变。"""
    from docx import Document
    from vermes_cli.blueprints import artifacts as art

    path = f"/tmp/{_tmp_name('.docx')}"
    try:
        doc = Document()
        doc.add_paragraph("段落零 不变")
        doc.add_paragraph("段落一 待改")
        doc.add_paragraph("段落二 不变")
        doc.save(path)

        before = [p.text for p in Document(path).paragraphs]
        assert before == ["段落零 不变", "段落一 待改", "段落二 不变"]

        n = art._apply_docx_paragraphs(path, [{"i": 1, "text": "段落一 已改"}])
        assert n == 1

        after = [p.text for p in Document(path).paragraphs]
        assert after == ["段落零 不变", "段落一 已改", "段落二 不变"]
    finally:
        if os.path.exists(path):
            os.unlink(path)


@pytest.mark.skipif(not _pdf_engine_available(), reason="本机无可用 pandoc PDF 引擎（需 LaTeX 或 weasyprint）")
def test_regenerate_pdf_from_md():
    """真实 pandoc：从 md 重生成 pdf，产物以 %PDF 开头。"""
    from vermes_cli.blueprints import artifacts as art

    md_path = f"/tmp/{_tmp_name('.md')}"
    pdf_path = f"/tmp/{_tmp_name('.pdf')}"
    try:
        with open(md_path, "w", encoding="utf-8") as f:
            f.write("# 标题\n\n正文内容。\n")
        open(pdf_path, "w").close()  # 端点要求文件已存在
        art._regenerate_pdf_from_md(pdf_path, "# 标题\n\n正文内容重生成。\n")
        with open(pdf_path, "rb") as f:
            assert f.read(5).startswith(b"%PDF")
    finally:
        for p in (md_path, pdf_path):
            if os.path.exists(p):
                os.unlink(p)


def test_editable_and_write_endpoints_via_app():
    """端点级：GET /editable（docx/pdf）+ POST /content（docx JSON / pdf md / xlsx 字节）。"""
    from docx import Document
    import vermes_cli.web_server as ws
    from starlette.testclient import TestClient

    client = TestClient(ws.app)

    # ── docx ──
    docx_path = f"/tmp/{_tmp_name('.docx')}"
    try:
        doc = Document()
        doc.add_paragraph("摘要：原始。")
        doc.add_paragraph("引言：原始。")
        doc.save(docx_path)

        rel = "tmp/" + os.path.basename(docx_path)
        r = client.get(f"/api/v1/artifacts/{rel}/editable")
        assert r.status_code == 200, r.text
        body = r.json()
        assert body["type"] == "docx"
        assert body["paragraphs"][0]["text"] == "摘要：原始。"

        payload = {"type": "docx", "paragraphs": [{"i": 0, "text": "摘要：已编辑。"}]}
        w = client.post(f"/api/v1/artifacts/{rel}/content",
                        json=payload, headers={"Content-Type": "application/json"})
        assert w.status_code == 200, w.text
        assert w.json()["updated"] == 1

        r2 = client.get(f"/api/v1/artifacts/{rel}/editable")
        assert r2.json()["paragraphs"][0]["text"] == "摘要：已编辑。"
    finally:
        if os.path.exists(docx_path):
            os.unlink(docx_path)

    # ── pdf（无引擎则跳过该子节，机制待用户环境验证）──
    if _pdf_engine_available():
        md_path = f"/tmp/{_tmp_name('.md')}"
        pdf_path = f"/tmp/{_tmp_name('.pdf')}"
        try:
            with open(md_path, "w", encoding="utf-8") as f:
                f.write("# PDF 源\n")
            open(pdf_path, "w").close()
            rel = "tmp/" + os.path.basename(pdf_path)
            r = client.get(f"/api/v1/artifacts/{rel}/editable")
            assert r.status_code == 200
            assert r.json()["regenerable"] is True
            assert "PDF 源" in r.json()["source_md"]

            w = client.post(f"/api/v1/artifacts/{rel}/content",
                            json={"type": "pdf", "md": "# PDF 重生成\n"},
                            headers={"Content-Type": "application/json"})
            assert w.status_code == 200, w.text
            with open(pdf_path, "rb") as f:
                assert f.read(5).startswith(b"%PDF")
        finally:
            for p in (md_path, pdf_path):
                if os.path.exists(p):
                    os.unlink(p)
    else:
        # pdf 子节跳过（无 pandoc PDF 引擎）；不调 pytest.skip() 以免终止整个测试函数
        # 机制待用户环境（MacTeX / weasyprint）验证
        pass

    # ── xlsx 二进制回存冒烟（方案 A：前端 SheetJS 重生成后 POST 字节）──
    xlsx_path = f"/tmp/{_tmp_name('.xlsx')}"
    try:
        import zipfile
        with zipfile.ZipFile(xlsx_path, "w") as z:
            z.writestr("[Content_Types].xml", "<Types/>")
        rel = "tmp/" + os.path.basename(xlsx_path)
        with open(xlsx_path, "rb") as f:
            blob = f.read()
        w = client.post(f"/api/v1/artifacts/{rel}/content",
                        content=blob, headers={"Content-Type": "application/octet-stream"})
        assert w.status_code == 200, w.text
        assert w.json()["ok"] is True
    finally:
        if os.path.exists(xlsx_path):
            os.unlink(xlsx_path)
